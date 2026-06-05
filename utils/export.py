"""
export.py — Exporta a peça gerada para .docx.
Usa python-docx com formatação jurídica básica.

Este módulo foi movido de export.py raiz para utils/export.py.
"""
import io
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False


def _aplicar_estilo_juridico(doc: "Document") -> None:
    """Configura margens e fonte padrão para petições."""
    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    # Estilo normal
    estilo = doc.styles["Normal"]
    font = estilo.font
    font.name = "Arial"
    font.size = Pt(12)

    paragrafo = estilo.paragraph_format
    paragrafo.space_after = Pt(6)
    paragrafo.line_spacing = Pt(18)  # 1.5


def _adicionar_paragrafo(doc: "Document", texto: str, negrito: bool = False, centralizado: bool = False) -> None:
    p = doc.add_paragraph()
    if centralizado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = negrito
    run.font.name = "Arial"
    run.font.size = Pt(12)


def gerar_docx(peca_texto: str, codigo: str, nome_arquivo: str = "") -> bytes:
    """
    Converte o texto Markdown da peça em .docx jurídico.
    Retorna bytes prontos para st.download_button.
    """
    if not DOCX_DISPONIVEL:
        raise ImportError("python-docx não instalado. Execute: pip install python-docx")

    doc = Document()
    _aplicar_estilo_juridico(doc)

    linhas = peca_texto.split("\n")
    em_checklist = False

    for linha in linhas:
        linha_strip = linha.strip()

        # Ignora checklist e pendências no .docx (são metadados)
        if linha_strip.startswith("## CHECKLIST") or linha_strip.startswith("## PENDÊNCIAS"):
            em_checklist = True
            continue
        if em_checklist and linha_strip.startswith("## ") and not linha_strip.startswith("## CHECKLIST") and not linha_strip.startswith("## PENDÊNCIAS"):
            em_checklist = False
        if em_checklist:
            continue

        # Cabeçalhos markdown → negrito centralizado
        if linha_strip.startswith("# "):
            _adicionar_paragrafo(doc, linha_strip[2:], negrito=True, centralizado=True)
        elif linha_strip.startswith("## "):
            _adicionar_paragrafo(doc, linha_strip[3:], negrito=True, centralizado=True)
        elif linha_strip.startswith("### "):
            _adicionar_paragrafo(doc, linha_strip[4:], negrito=True)
        elif linha_strip.startswith("**") and linha_strip.endswith("**"):
            # Linha inteiramente em negrito
            _adicionar_paragrafo(doc, linha_strip.strip("*"), negrito=True)
        elif linha_strip == "" or linha_strip == "---":
            doc.add_paragraph()
        else:
            # Remove markdown inline restante
            texto_limpo = re.sub(r"\*\*(.*?)\*\*", r"\1", linha_strip)
            texto_limpo = re.sub(r"\*(.*?)\*",   r"\1", texto_limpo)
            texto_limpo = re.sub(r"`(.*?)`",     r"\1", texto_limpo)
            _adicionar_paragrafo(doc, texto_limpo)

    # Rodapé com data de geração
    doc.add_paragraph()
    _adicionar_paragrafo(
        doc,
        f"[Gerado pelo Sistema Jurídico IA em {datetime.now().strftime('%d/%m/%Y às %H:%M')}]",
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def nome_arquivo_peca(codigo: str, autor: str = "") -> str:
    data = datetime.now().strftime("%Y%m%d")
    autor_slug = autor.split()[0].lower() if autor else "parte"
    return f"{codigo}_{autor_slug}_{data}.docx"
